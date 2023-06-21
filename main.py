import shutil

import docx
from docx2pdf import convert
import os
from tkinter import *
from tkinter import filedialog as fd
from tkinter import messagebox as mb

shablon_name = "xxxx"
studentnames_list = "xxxx"
doc = None


def create_file(document, doc_name):
    doc.save(f'certificates/{doc_name}.docx')
    convert(f"certificates/{doc_name}.docx")
    os.remove(f'certificates/{doc_name}.docx')


def check_tables():
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text == "??":
                    names = open(studentnames_list, encoding='utf-8')
                    counter = 0
                    for name in names:
                        name = name.strip()
                        cell.text = name
                        print(name)
                        create_file(doc, name)
                        counter += 1
                        print(f"Созданы {counter} грамоты")
                    names.close()


def check_runs():
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text == "??":
                names = open(studentnames_list, encoding='utf-8')
                counter = 0
                for name in names:
                    name = name.strip()
                    run = run.clear()
                    print(name)
                    run.add_text(name)
                    create_file(doc, name)
                    counter += 1
                print(f"Созданы {counter} грамоты")
                names.close()


def choose_doc():
    global shablon_name
    global shablon_way
    shablon_name = fd.askopenfilename(
        filetypes=(("DOCX files", "*.docx"),))
    shablon_way.configure(text=shablon_name)


window = Tk()
window.title("Генератор сертификатов")
window.geometry('400x400')
frame = Frame(window,
              padx=15,
              pady=15
              )
frame.pack(expand=True)
open_shablon = Button(frame, text="Выбрать шаблон", command=choose_doc)
open_shablon.grid(row=1)
shablon_way = Label(
    frame,
    text="Путь"
)
shablon_way.grid(row=3)


def choose_list():
    global studentnames_list
    global list_way
    studentnames_list = fd.askopenfilename(
        filetypes=(("TXT files", "*.txt"),))
    list_way.configure(text=studentnames_list)


open_list = Button(frame, text="Выбрать список", command=choose_list)
open_list.grid(row=4)
list_way = Label(
    frame,
    text="Путь"
)
list_way.grid(row=6)


def generate():
    global doc
    doc = docx.Document(shablon_name)
    if os.path.exists("certificates"):
        shutil.rmtree("certificates")
    os.mkdir("certificates")
    check_tables()
    check_runs()
    mb.showinfo(title="Успех", message="Готово!")


open_list = Button(frame, text="Генерировать", command=generate)
open_list.grid(row=8)

if os.path.exists("certificates"):
    shutil.rmtree("certificates")
os.mkdir("certificates")

window.mainloop()
print(shablon_name)
